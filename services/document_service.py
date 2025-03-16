# services/document_service.py
import os
import logging
import subprocess
from PIL import Image
import io

logger = logging.getLogger(__name__)

class DocumentService:
    """Service for document conversion operations."""
    
    def __init__(self, temp_dir='temp'):
        """Initialize document service with temporary directory for processing."""
        self.temp_dir = temp_dir
        os.makedirs(temp_dir, exist_ok=True)
    
    def convert_document(self, input_path, output_path):
        """
        Convert document using appropriate library based on input/output format.
        
        Args:
            input_path (str): Path to the input document file
            output_path (str): Path where the converted document will be saved
        """
        input_format = input_path.split('.')[-1].lower()
        output_format = output_path.split('.')[-1].lower()
        
        logger.info(f"Converting document from {input_format} to {output_format}")
        
        try:
            if output_format == 'pdf':
                self._convert_to_pdf(input_path, output_path, input_format)
            elif output_format == 'docx':
                self._convert_to_docx(input_path, output_path, input_format)
            elif output_format in ('txt', 'md', 'html'):
                self._convert_to_text_format(input_path, output_path, input_format, output_format)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")
                
        except Exception as e:
            logger.error(f"Document conversion error: {str(e)}")
            raise
    
    def _convert_to_pdf(self, input_path, output_path, input_format):
        """Convert a document to PDF format."""
        if input_format == 'docx':
            # Convert DOCX to PDF
            from docx2pdf import convert
            convert(input_path, output_path)
        
        elif input_format in ('txt', 'md', 'html'):
            # Convert text-based document to PDF
            import pdfkit
            
            # Read the input file
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Process content for markdown
            if input_format == 'md':
                import markdown
                content = markdown.markdown(content)
            
            # Use pdfkit for conversion
            options = {
                'encoding': 'UTF-8',
                'page-size': 'A4',
                'margin-top': '20mm',
                'margin-right': '20mm',
                'margin-bottom': '20mm',
                'margin-left': '20mm',
            }
            pdfkit.from_string(content, output_path, options=options)
        
        else:
            raise ValueError(f"Conversion from {input_format} to PDF is not supported")
    
    def _convert_to_docx(self, input_path, output_path, input_format):
        """Convert a document to DOCX format."""
        from docx import Document
        
        if input_format == 'pdf':
            # Use pdf2docx for PDF to DOCX conversion
            from pdf2docx import Converter
            cv = Converter(input_path)
            cv.convert(output_path)
            cv.close()
        
        elif input_format in ('txt', 'md', 'html'):
            # Create a new document
            doc = Document()
            
            # Read the input file
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Process content based on format
            if input_format == 'md':
                import markdown
                from bs4 import BeautifulSoup
                
                # Convert markdown to HTML, then extract text
                html_content = markdown.markdown(content)
                soup = BeautifulSoup(html_content, 'html.parser')
                content = soup.get_text()
                
                # Add headings and paragraphs
                lines = content.split('\n')
                for line in lines:
                    if line.strip():
                        if line.startswith('# '):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith('## '):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith('### '):
                            doc.add_heading(line[4:], level=3)
                        else:
                            doc.add_paragraph(line)
            
            elif input_format == 'html':
                from bs4 import BeautifulSoup
                
                # Parse HTML and extract text
                soup = BeautifulSoup(content, 'html.parser')
                
                # Process headings
                for i in range(1, 7):
                    for heading in soup.find_all(f'h{i}'):
                        doc.add_heading(heading.get_text(), level=i)
                        heading.decompose()
                
                # Process paragraphs
                for paragraph in soup.find_all('p'):
                    doc.add_paragraph(paragraph.get_text())
            
            else:  # Plain text
                # Add each line as a paragraph
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
            
            # Save the document
            doc.save(output_path)
        
        else:
            raise ValueError(f"Conversion from {input_format} to DOCX is not supported")
    
    def _convert_to_text_format(self, input_path, output_path, input_format, output_format):
        """Convert a document to TXT, MD, or HTML format."""
        if input_format == 'pdf':
            # Extract text from PDF
            from pdfminer.high_level import extract_text
            text = extract_text(input_path)
            
            if output_format == 'html':
                # Convert to basic HTML
                html_content = f"<html><body><pre>{text}</pre></body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            
            elif output_format == 'md':
                # Convert to basic Markdown
                # For PDF to MD, just use the plain text
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            
            else:  # txt
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
        
        elif input_format == 'docx':
            # Extract text from DOCX
            from docx import Document
            doc = Document(input_path)
            
            # Extract paragraphs
            text = '\n\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            if output_format == 'html':
                # Convert to basic HTML
                paragraphs_html = ''
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs_html += f"<p>{paragraph.text}</p>\n"
                
                html_content = f"<html><body>\n{paragraphs_html}</body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            
            elif output_format == 'md':
                # Convert to basic Markdown
                with open(output_path, 'w', encoding='utf-8') as f:
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            # Check if it's a heading based on style
                            if paragraph.style.name.startswith('Heading'):
                                level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                                f.write('#' * level + ' ' + paragraph.text + '\n\n')
                            else:
                                f.write(paragraph.text + '\n\n')
            
            else:  # txt
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
        
        elif input_format == 'html':
            # Process HTML
            from bs4 import BeautifulSoup
            
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            if output_format == 'md':
                # Convert HTML to Markdown
                import html2text
                h = html2text.HTML2Text()
                h.ignore_links = False
                markdown_text = h.handle(content)
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_text)
            
            elif output_format == 'txt':
                # Extract plain text
                text = soup.get_text()
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            
            else:  # html to html, just copy
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
        
        elif input_format == 'md':
            # Process Markdown
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if output_format == 'html':
                # Convert Markdown to HTML
                import markdown
                html_content = markdown.markdown(content)
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(f"<html><body>\n{html_content}\n</body></html>")
            
            elif output_format == 'txt':
                # For MD to TXT, just remove Markdown syntax (simplified approach)
                import re
                
                # Remove headings
                text = re.sub(r'^#+\s+', '', content, flags=re.MULTILINE)
                
                # Remove emphasis
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
                text = re.sub(r'\*(.+?)\*', r'\1', text)      # Italic
                text = re.sub(r'__(.+?)__', r'\1', text)      # Bold
                text = re.sub(r'_(.+?)_', r'\1', text)        # Italic
                
                # Remove links
                text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            
            else:  # md to md, just copy
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
        
        elif input_format == 'txt':
            # Process plain text
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if output_format == 'html':
                # Convert TXT to basic HTML
                html_content = "<html><body>"
                for line in content.split('\n'):
                    if line.strip():
                        html_content += f"<p>{line}</p>\n"
                html_content += "</body></html>"
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            
            elif output_format == 'md':
                # For TXT to MD, just copy (no transformation needed)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            else:  # txt to txt, just copy
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
        
        else:
            raise ValueError(f"Conversion from {input_format} to {output_format} is not supported")